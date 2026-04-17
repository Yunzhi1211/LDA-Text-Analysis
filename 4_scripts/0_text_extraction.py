"""
LDA主题分析 - 文本提取模块
从Word文档批量提取文本内容
"""

import os
from pathlib import Path
from docx import Document
import logging

logger = logging.getLogger(__name__)


class TextExtractor:
    """从DOCX文件提取文本的类"""
    
    def __init__(self, input_dir: str):
        self.input_dir = Path(input_dir)
        if not self.input_dir.exists():
            raise FileNotFoundError(f"输入文件夹不存在: {self.input_dir}")
    
    def extract_from_docx(self, filepath: str) -> str:
        try:
            doc = Document(filepath)
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)
            full_text = "\n".join(texts)
            return full_text
        except Exception as e:
            logger.error(f"提取文件 {filepath} 失败: {e}")
            return ""
    
    def extract_all(self) -> dict:
        documents = {}
        docx_files = list(self.input_dir.glob("*.docx"))
        
        if not docx_files:
            logger.warning(f"在 {self.input_dir} 中未找到.docx文件")
            return documents
        
        logger.info(f"找到 {len(docx_files)} 个DOCX文件，开始提取...")
        
        for idx, filepath in enumerate(docx_files, 1):
            filename = filepath.name
            text = self.extract_from_docx(str(filepath))
            documents[filename] = text
            logger.info(f"[{idx}/{len(docx_files)}] 已提取: {filename}")
        
        return documents


def extract_texts(input_dir: str) -> dict:
    extractor = TextExtractor(input_dir)
    return extractor.extract_all()
